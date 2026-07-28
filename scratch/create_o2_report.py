import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

doc = Document()

# Page Setup - Margins 1 inch
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Helper for shading table cells
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('รายงานสรุปวิเคราะห์การใช้งานออกซิเจนทางการแพทย์\nโรงพยาบาลนครพิงค์ และศูนย์มะเร็งขอนตาล')
run.font.name = 'TH Sarabun PSK'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = subtitle_p.add_run('ข้อมูลสถิติประวัติการใช้ 31 เดือน (กันยายน 2566 – มีนาคม 2569) และข้อมูลโทรมาตร Real-time\nกลุ่มงานวิศวกรรมความปลอดภัยและระบบก๊าซทางการแพทย์ โรงพยาบาลนครพิงค์')
srun.font.name = 'TH Sarabun PSK'
srun.font.size = Pt(14)
srun.font.italic = True
srun.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Executive Summary Section
h1 = doc.add_paragraph()
hrun = h1.add_run('1. สรุปสาระสำคัญสำหรับบริหาร (Executive Summary)')
hrun.font.name = 'TH Sarabun PSK'
hrun.font.size = Pt(16)
hrun.font.bold = True
hrun.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.15
prun = p.add_run(
    'ระบบออกซิเจนทางการแพทย์เป็นระบบสาธารณูปโภคขั้นวิกฤต (Critical Healthcare Infrastructure) ของโรงพยาบาลนครพิงค์ '
    'จากข้อมูลการบันทึกปริมาณการใช้งานออกซิเจนเหลวสะสมจำนวน 31 เดือน (ตั้งแต่ ก.ย. 2566 ถึง มี.ค. 2569) พบว่า อัตราการใช้ออกซิเจนเหลวเฉลี่ยรวมขององค์กรอยู่ที่ '
)
prun.font.name = 'TH Sarabun PSK'
prun.font.size = Pt(14)

b1 = p.add_run('69,054.24 ลูกบาศก์เมตรต่อเดือน (m³/เดือน)')
b1.font.name = 'TH Sarabun PSK'
b1.font.size = Pt(14)
b1.font.bold = True

prun2 = p.add_run(
    ' โดยแบ่งเป็นปริมาณการใช้ของโรงพยาบาลนครพิงค์ (ถังหลัก 37,000 ลิตร) สัดส่วน 97.28% (เฉลี่ย 67,173.09 m³/เดือน) '
    'และศูนย์มะเร็งขอนตาล สัดส่วน 2.72% (เฉลี่ย 2,159.84 m³/เดือน)\n\n'
    'ในช่วงไตรมาสแรกของปี 2569 ปริมาณการใช้ออกซิเจนเหลวปรับตัวสูงขึ้นอย่างต่อเนื่อง โดยทำสถิติสูงสุดใหม่ (Peak Consumption) '
    'ในเดือนมกราคม 2569 รวมสูงถึง 80,182.24 m³ อันเนื่องมาจากการเพิ่มขึ้นของผู้ป่วยวิกฤต และการเปิดให้บริการพื้นที่รักษาใหม่ '
    'ทั้งนี้ ถังออกซิเจนเหลวหลักและระบบจ่ายแก๊สสำรองประจำอาคารได้รับการเฝ้าระวังผ่านระบบโทรมาตรทางไกล (IoT Telemetry) '
    'และมีแผนสำรองท่อแก๊สที่สามารถรองรับการจ่ายแก๊สต่อเนื่องได้ 9.44 ชั่วโมง ซึ่งเพียงพอต่อระยะเวลาการจัดส่งฉุกเฉินภายใน 3 ชั่วโมง'
)
prun2.font.name = 'TH Sarabun PSK'
prun2.font.size = Pt(14)

doc.save(r'E:\000_Antigraviti\MedicalGasNKP\Oxygen\07_Liquid_Oxygen_Main_System\รายงานสรุปการใช้ออกซิเจน.docx')
print('Initial Word Doc created!')
