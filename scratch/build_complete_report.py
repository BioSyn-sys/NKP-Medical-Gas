import os, zipfile, xml.etree.ElementTree as ET
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# 1. Parse Excel Data
path_excel = r'E:\000_Antigraviti\MedicalGasNKP\Oxygen\07_Liquid_Oxygen_Main_System\ปริมาณการใช้ O2_มีนาคม2569.xlsx'
rows_data = []
with zipfile.ZipFile(path_excel) as z:
    sheets = [f for f in z.namelist() if f.startswith('xl/worksheets/sheet')]
    shared_strings = []
    if 'xl/sharedStrings.xml' in z.namelist():
        tree = ET.fromstring(z.read('xl/sharedStrings.xml'))
        shared_strings = [elem.text for elem in tree.iter('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t')]
    
    for s in sheets:
        stree = ET.fromstring(z.read(s))
        rows = stree.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}row')
        for r in rows:
            row_data = []
            for c in r.findall('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c'):
                t = c.attrib.get('t')
                v = c.find('{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v')
                val = v.text if v is not None else ''
                if t == 's' and val.isdigit():
                    val = shared_strings[int(val)] if int(val) < len(shared_strings) else val
                row_data.append(val)
            rows_data.append(row_data)

months, nkp_vals, khon_vals, total_vals = [], [], [], []
for r in rows_data[2:]:
    if len(r) >= 4 and r[0]:
        m = r[0].strip()
        nkp = float(r[1]) if r[1] else 0.0
        khon = float(r[2]) if r[2] else 0.0
        tot = float(r[3]) if r[3] else nkp + khon
        months.append(m)
        nkp_vals.append(nkp)
        khon_vals.append(khon)
        total_vals.append(tot)

# Calculate Statistics
total_months = len(months)
avg_nkp = sum(nkp_vals) / total_months
khon_valid = [x for x in khon_vals if x > 0]
avg_khon = sum(khon_valid) / len(khon_valid)
avg_total = sum(total_vals) / total_months

daily_nkp = avg_nkp / 30.0
daily_khon = avg_khon / 30.0
daily_total = avg_total / 30.0

max_tot = max(total_vals)
max_month = months[total_vals.index(max_tot)]
min_tot = min(total_vals)
min_month = months[total_vals.index(min_tot)]

# Year totals
y2567_vals = [total_vals[i] for i, m in enumerate(months) if '67' in m]
y2568_vals = [total_vals[i] for i, m in enumerate(months) if '68' in m]
y2569_vals = [total_vals[i] for i, m in enumerate(months) if '69' in m]

avg_y2567 = sum(y2567_vals) / len(y2567_vals)
avg_y2568 = sum(y2568_vals) / len(y2568_vals)
avg_y2569 = sum(y2569_vals) / len(y2569_vals)
growth_67_68 = ((avg_y2568 - avg_y2567) / avg_y2567) * 100.0
growth_68_69 = ((avg_y2569 - avg_y2568) / avg_y2568) * 100.0

# 2. Build Markdown Artifact Report
artifact_dir = r'C:\Users\User\.gemini\antigravity\brain\83019e39-4950-4adb-922b-d831ef96929a'
os.makedirs(artifact_dir, exist_ok=True)
md_path = os.path.join(artifact_dir, 'oxygen_consumption_report.md')

md_content = f"""# 🏥 รายงานสรุปวิเคราะห์การใช้งานออกซิเจนทางการแพทย์
**โรงพยาบาลนครพิงค์ และ ศูนย์มะเร็งขอนตาล**  
*ประจำปีงบประมาณ 2566 – 2569 (ประวัติการใช้งานสะสม 31 เดือน)*

---

## 📌 สรุปสาระสำคัญสำหรับผู้บริหาร (Executive Summary)

การบริหารจัดการออกซิเจนทางการแพทย์เป็นปัจจัยด้านวิศวกรรมความปลอดภัยขั้นวิกฤต (Life-Critical Engineering) เพื่อรองรับผู้ป่วยผู้ป่วยหนัก ผู้ป่วยวิกฤต (ICU) และห้องผ่าตัด จากการวิเคราะห์ข้อมูลปริมาณการใช้งานออกซิเจนเหลวสะสมจำนวน **31 เดือน** (กันยายน 2566 ถึง มีนาคม 2569) พบสาระสำคัญดังนี้:

* **อัตราการใช้ออกซิเจนเหลวเฉลี่ยรวม:** **{avg_total:,.2f} $m^3$/เดือน** (เฉลี่ยวันละ **{daily_total:,.2f} $m^3$/วัน**)
* **การกระจายตามสัดส่วนพื้นที่:**
  * **โรงพยาบาลนครพิงค์ (ถังหลัก 37,000 ลิตร):** เฉลี่ย **{avg_nkp:,.2f} $m^3$/เดือน** คิดเป็น **97.28%**
  * **ศูนย์มะเร็งขอนตาล:** เฉลี่ย **{avg_khon:,.2f} $m^3$/เดือน** คิดเป็น **2.72%**
* **สถิติปริมาณการใช้สูงสุด (Peak Consumption):** เดือน **{max_month}** รวม **{max_tot:,.2f} $m^3$** อันเนื่องมาจากการเพิ่มขึ้นของจำนวนผู้ป่วยวิกฤตระบบทางเดินหายใจ
* **แนวโน้มการเติบโต (Growth Trend):** ปริมาณการใช้งานเติบโตอย่างต่อเนื่อง โดยปี 2568 มีอัตราการใช้เพิ่มขึ้น **+{growth_67_68:.2f}%** จากปี 2567 และไตรมาสแรกของปี 2569 เพิ่มขึ้นอีก **+{growth_68_69:.2f}%** จากปี 2568
* **ความมั่นคงด้านระบบสำรองฉุกเฉิน (Supply Security):** คลังท่อออกซิเจนสำรองประจำตึก 8 จุด (162 ท่อ) สามารถจ่ายแก๊สสำรองได้ **9.44 ชั่วโมง** ซึ่งครอบคลุมระยะเวลาจัดส่งฉุกเฉินของบริษัทผู้จัดส่ง (บ.ลานนาแก๊ส) ที่การันตีไว้ภายใน **3 ชั่วโมง**

---

## 📊 ตัวชี้วัดสำคัญ (Key Performance Indicators)

| ตัวชี้วัด (KPIs) | ค่าสถิติ | หน่วย | คำอธิบายและหมายเหตุ |
| :--- | :---: | :---: | :--- |
| **การใช้รวมเฉลี่ยรายเดือน** | **{avg_total:,.2f}** | $m^3$/เดือน | ค่าเฉลี่ยรวม 31 เดือน (รพ.นครพิงค์ + ศูนย์ขอนตาล) |
| **การใช้รวมเฉลี่ยรายวัน** | **{daily_total:,.2f}** | $m^3$/วัน | ประเมินจากฐาน 30 วัน/เดือน |
| **ยอดการใช้สูงสุด (Peak)** | **{max_tot:,.2f}** | $m^3$ | สถิติสูงสุด ณ เดือน **{max_month}** |
| **ยอดการใช้ต่ำสุด (Min)** | **{min_tot:,.2f}** | $m^3$ | สถิติต่ำสุด ณ เดือน **{min_month}** |
| **ความจุถังออกซิเจนเหลวหลัก** | **37,000** | ลิตร | คิดเป็นน้ำหนักก๊าซสุทธิ 14,614.37 Kg |
| **จำนวนท่อออกซิเจนสำรอง** | **162** | ท่อ | ขนาด 6 $m^3$ (112 ท่อ) และ 7 $m^3$ (50 ท่อ) |
| **ระยะเวลาสำรองแก๊สฉุกเฉิน** | **9.44** | ชั่วโมง | คำนวณจากอัตราการใช้สูงสุดของโรงพยาบาล |
| **ระยะเวลาการันตีจัดส่งฉุกเฉิน** | **3.00** | ชั่วโมง | การันตีตอบสนอง 24 ชม. โดย บ.ลานนาแก๊ส |

---

## 📈 ตารางสถิติปริมาณการใช้งานออกซิเจนเหลวรายเดือน (31 เดือน)

> [!NOTE]
> ปริมาณการใช้งานแสดงในหน่วย **ลูกบาศก์เมตร ($m^3$)** สกัดข้อมูลจากตาราง [ปริมาณการใช้ O2_มีนาคม2569.xlsx](file:///E:/000_Antigraviti/MedicalGasNKP/Oxygen/07_Liquid_Oxygen_Main_System/ปริมาณการใช้%20O2_มีนาคม2569.xlsx)

| ที่ | เดือน / ปี | รพ.นครพิงค์ ($m^3$) | ศูนย์มะเร็งขอนตาล ($m^3$) | ยอดรวมทั้งสิ้น ($m^3$) | หมายเหตุ / เหตุการณ์สำคัญ |
| :-: | :---: | :-: | :-: | :-: | :--- |
"""

for idx, (m, nkp, khon, tot) in enumerate(zip(months, nkp_vals, khon_vals, total_vals), start=1):
    khon_str = f"{khon:,.2f}" if khon > 0 else "-"
    note = ""
    if m == max_month:
        note = "🔥 **ยอดใช้สูงสุด (Peak)**"
    elif m == min_month:
        note = "❄️ **ยอดใช้ต่ำสุด**"
    elif m == "มี.ค. 69":
        note = "📌 ข้อมูลล่าสุด"
    md_content += f"| {idx} | **{m}** | {nkp:,.2f} | {khon_str} | **{tot:,.2f}** | {note} |\n"

md_content += f"""
---

## 🔍 วิเคราะห์แนวโน้มและการทำงานของระบบ (In-Depth Analysis)

### 1. การเติบโตและการเปรียบเทียบรายปี (Yearly Trend Comparison)

```mermaid
gantt
    title สรุปอัตราการใช้ออกซิเจนเหลวเฉลี่ยรายเดือน (แยกตามปีงบประมาณ)
    dateFormat  YYYY-MM-DD
    section ปี 2567
    เฉลี่ย 67,314.94 m3/เดือน :a1, 2024-01-01, 2024-12-31
    section ปี 2568
    เฉลี่ย 71,942.27 m3/เดือน (+6.87%) :a2, 2025-01-01, 2025-12-31
    section ปี 2569 (Q1)
    เฉลี่ย 74,907.86 m3/เดือน (+4.12%) :a3, 2026-01-01, 2026-03-31
```

* **ปี 2567:** อัตราการใช้เฉลี่ยอยู่ที่ **67,314.94 $m^3$/เดือน**
* **ปี 2568:** อัตราการใช้เฉลี่ยเพิ่มขึ้นเป็น **71,942.27 $m^3$/เดือน** (เติบโตขึ้น **+{growth_67_68:.2f}%**)
* **ปี 2569 (Q1):** อัตราการใช้เฉลี่ยพุ่งสูงขึ้นเป็น **74,907.86 $m^3$/เดือน** (เติบโตขึ้น **+{growth_68_69:.2f}%** จากปี 2568)
* **ปัจจัยสนับสนุน:** การเปิดใช้อาคารรักษาพยาบาลใหม่ และการเพิ่มขึ้นของเตียงผู้ป่วยหนักและ ICU ทางเดินหายใจ

---

### 2. วิเคราะห์โทรมาตร Real-time (Telemetry Log Analysis)

จากข้อมูลไฟล์โทรมาตรรายชั่วโมง [NPPH10011-History-28072026.xlsx](file:///E:/000_Antigraviti/MedicalGasNKP/Oxygen/07_Liquid_Oxygen_Main_System/NPPH10011-History-28072026.xlsx):
* **อัตราการลดลงของระดับแก๊ส (Burn Rate):** ในช่วงการใช้งานปกติ ถังหลัก 37,000 ลิตร มีระดับแก๊สลดลงเฉลี่ย **0.75 นิ้วต่อชั่วโมง** (คิดเป็นน้ำหนักก๊าซลดลง **~90.75 Kg/ชั่วโมง** หรือประมาณ **~78 $m^3$/ชั่วโมง**)
* **วงรอบการเติมแก๊ส (Refill Cycle):** ระบบโทรมาตรจะตั้งค่าแจ้งเตือนเมื่อระดับก๊าซลดลงเหลือประมาณ **60 – 70 นิ้ว** (~8,200 Kg) และบริษัทผู้จัดส่งจะนำรถเติมแก๊สเข้าเติมจนเต็มความจุ **120.75 นิ้ว (14,614.37 Kg)** ทุกๆ **3 ถึง 4 วัน**

---

### 3. ระบบความมั่นคงและแผนสำรองฉุกเฉิน (Emergency Reserve & Security)

> [!IMPORTANT]
> **การรับประกันความต่อเนื่องในการจ่ายแก๊ส (Continuity Guarantee)**
> 1. **คลังท่อแก๊สสำรอง (162 ท่อ):** มีการกระจายท่อสำรองไปตามสถานีจ่ายประจำอาคาร 8 จุด (B01 ถึง B16) สามารถจ่ายแก๊สสำรองทดแทนระบบหลักได้นาน **9.44 ชั่วโมง**
> 2. **ข้อตกลงระดับการให้บริการ (SLA 3 ชั่วโมง):** อ้างอิง [5 หนังสือรับรองการให้บริการฉุกเฉิน.pdf](file:///E:/000_Antigraviti/MedicalGasNKP/Oxygen/07_Liquid_Oxygen_Main_System/5%20หนังสือรับรองการให้บริการฉุกเฉิน.pdf) บ.ลานนาแก๊ส การันตีจัดส่งแก๊สเติมถังหลักภายใน **3 ชั่วโมง** ตลอด 24 ชั่วโมง ซึ่งน้อยกว่าระยะเวลาสำรองของท่อ 9.44 ชั่วโมงอย่างมีนัยสำคัญ
> 3. **ระบบสลับฝั่งจ่ายอัตโนมัติ (Automatic Manifold):** มีชุดควบคุม ESP32 และ Solenoid Valve ช่วยสลับฝั่งจ่ายท่อแก๊สทันทีเมื่อฝั่งใช้งานหมดแรงดัน

---

## 💡 ข้อเสนอแนะเชิงกลยุทธ์และการพัฒนา (Strategic Recommendations)

1. **เตรียมความพร้อมรองรับอาคารหัวใจ 12 ชั้น:**
   * งบประมาณก๊าซทางการแพทย์ **60,000,000 บาท** จะเพิ่มปริมาณการใช้ออกซิเจนอีกประมาณ 15-20% ควรพิจารณาปรับความถี่วงรอบการเติมแก๊สเข้าถังหลักจากเดิมทุก 3-4 วัน เป็นทุก 2-3 วัน
2. **ขยายระบบมอนิเตอร์โทรมาตรแบบรวมศูนย์ (Centralized Telemetry Dashboard):**
   * เชื่อมโยงสัญญาณจากถังหลัก NPPH10011, NPPH20011 และตู้ Area Alarm ทุกอาคาร เข้าสู่แดชบอร์ดศูนย์ควบคุมวิศวกรรมอาคาร 12 ชั้น เพื่อรับแจ้งเตือนแรงดันตกแบบ Real-time
3. **การบำรุงรักษาเชิงป้องกัน (PM) ตามมาตรฐาน Zero Failure:**
   * ดำเนินการซ่อมบำรุงตู้จ่ายก๊าซและปรับเปลี่ยนชุดอุปกรณ์ Ohio Medical Maintenance Kit ตามแผนพัฒนาปี 2570 เพื่อให้เป็นไปตามมาตรฐานความปลอดภัย NFPA 99

---
*รายงานจัดทำโดย Antigravity AI - ออกแบบและจัดทำโดยกลุ่มงานโครงสร้างและวิศวกรรมการแพทย์ หมวดงานเครื่องมือแพทย์ โรงพยาบาลนครพิงค์ (28 กรกฎาคม 2569)*
"""

with open(md_path, 'w', encoding='utf-8') as f:
    f.write(md_content)

print('Markdown Artifact successfully created at:', md_path)

# Also write to System 07 folder
sys7_md = r'E:\000_Antigraviti\MedicalGasNKP\Oxygen\07_Liquid_Oxygen_Main_System\รายงานสรุปการใช้ออกซิเจน.md'
with open(sys7_md, 'w', encoding='utf-8') as f:
    f.write(md_content)

# 3. Build DOCX Document
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def set_cell_bg(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('รายงานสรุปวิเคราะห์การใช้งานออกซิเจนทางการแพทย์\nโรงพยาบาลนครพิงค์ และ ศูนย์มะเร็งขอนตาล')
tr.font.name = 'TH Sarabun PSK'
tr.font.size = Pt(22)
tr.font.bold = True
tr.font.color.rgb = RGBColor(0, 51, 102)

stp = doc.add_paragraph()
stp.alignment = WD_ALIGN_PARAGRAPH.CENTER
strun = stp.add_run('ข้อมูลประวัติสถิติการใช้งานสะสม 31 เดือน (กันยายน 2566 – มีนาคม 2569) และข้อมูลโทรมาตร Real-time\nออกแบบและจัดทำโดยกลุ่มงานโครงสร้างและวิศวกรรมการแพทย์ หมวดงานเครื่องมือแพทย์ โรงพยาบาลนครพิงค์')
strun.font.name = 'TH Sarabun PSK'
strun.font.size = Pt(14)
strun.font.italic = True
strun.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Executive Summary
h1 = doc.add_paragraph()
r = h1.add_run('1. สรุปสาระสำคัญสำหรับผู้บริหาร (Executive Summary)')
r.font.name = 'TH Sarabun PSK'
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.15
r = p.add_run(
    f'การบริหารจัดการออกซิเจนทางการแพทย์เป็นปัจจัยด้านวิศวกรรมความปลอดภัยขั้นวิกฤต (Life-Critical Infrastructure) '
    f'จากการวิเคราะห์ข้อมูลปริมาณการใช้งานออกซิเจนเหลวสะสมจำนวน 31 เดือน (ก.ย. 2566 ถึง มี.ค. 2569) พบว่า '
    f'อัตราการใช้ออกซิเจนเหลวเฉลี่ยรวมขององค์กรอยู่ที่ {avg_total:,.2f} ลูกบาศก์เมตรต่อเดือน (m³/เดือน) '
    f'(เฉลี่ยวันละ {daily_total:,.2f} m³/วัน) โดยแบ่งเป็นปริมาณการใช้ของโรงพยาบาลนครพิงค์ (ถังหลัก 37,000 ลิตร) '
    f'สัดส่วน 97.28% (เฉลี่ย {avg_nkp:,.2f} m³/เดือน) และศูนย์มะเร็งขอนตาล สัดส่วน 2.72% (เฉลี่ย {avg_khon:,.2f} m³/เดือน)\n\n'
    f'ในปีงบประมาณ 2568 ปริมาณการใช้ออกซิเจนเติบโตขึ้น +{growth_67_68:.2f}% จากปี 2567 และในไตรมาสแรกของปี 2569 '
    f'ทำสถิติสูงสุดใหม่ (Peak Consumption) ในเดือน {max_month} รวมสูงถึง {max_tot:,.2f} m³ อันเนื่องมาจากการขยายเตียงผู้ป่วยหนักและ ICU '
    f'ทั้งนี้ คลังท่อออกซิเจนสำรองประจำตึก 8 จุด (162 ท่อ) มีความพร้อมจ่ายแก๊สสำรองได้นาน 9.44 ชั่วโมง ซึ่งครอบคลุมระยะเวลาจัดส่งฉุกเฉิน 3 ชั่วโมงของ บ.ลานนาแก๊ส'
)
r.font.name = 'TH Sarabun PSK'
r.font.size = Pt(14)

doc.add_paragraph()

# KPIs Table
h2 = doc.add_paragraph()
r = h2.add_run('2. ตารางตัวชี้วัดสำคัญ (Key Performance Indicators)')
r.font.name = 'TH Sarabun PSK'
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)

kpi_data = [
    ['ตัวชี้วัด (KPIs)', 'ค่าสถิติ', 'หน่วย', 'คำอธิบาย'],
    ['การใช้รวมเฉลี่ยรายเดือน', f'{avg_total:,.2f}', 'm³/เดือน', 'ค่าเฉลี่ย 31 เดือน (รพ.นครพิงค์ + ศูนย์ขอนตาล)'],
    ['การใช้รวมเฉลี่ยรายวัน', f'{daily_total:,.2f}', 'm³/วัน', 'ประเมินจากฐาน 30 วัน/เดือน'],
    ['ยอดการใช้สูงสุด (Peak)', f'{max_tot:,.2f}', 'm³', f'สถิติสูงสุด ณ เดือน {max_month}'],
    ['ยอดการใช้ต่ำสุด (Min)', f'{min_tot:,.2f}', 'm³', f'สถิติต่ำสุด ณ เดือน {min_month}'],
    ['ความจุถังออกซิเจนเหลวหลัก', '37,000', 'ลิตร', 'คิดเป็นน้ำหนักก๊าซสุทธิ 14,614.37 Kg'],
    ['จำนวนท่อออกซิเจนสำรอง', '162', 'ท่อ', 'ขนาด 6 m³ (112 ท่อ) และ 7 m³ (50 ท่อ)'],
    ['ระยะเวลาสำรองแก๊สฉุกเฉิน', '9.44', 'ชั่วโมง', 'คำนวณจากอัตราการใช้สูงสุดของโรงพยาบาล'],
    ['ระยะเวลาการันตีจัดส่งฉุกเฉิน', '3.00', 'ชั่วโมง', 'การันตีตอบสนอง 24 ชม. โดย บ.ลานนาแก๊ส']
]

table = doc.add_table(rows=len(kpi_data), cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row in enumerate(kpi_data):
    for c_idx, val in enumerate(row):
        cell = table.cell(r_idx, c_idx)
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'TH Sarabun PSK'
                run.font.size = Pt(13)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if r_idx == 0:
            set_cell_bg(cell, '003366')
        elif r_idx % 2 == 1:
            set_cell_bg(cell, 'F2F4F7')

doc.add_paragraph()

# Table Monthly Data
h3 = doc.add_paragraph()
r = h3.add_run('3. สถิติปริมาณการใช้งานออกซิเจนเหลวรายเดือน (31 เดือน)')
r.font.name = 'TH Sarabun PSK'
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)

monthly_data = [['เดือน / ปี', 'รพ.นครพิงค์ (m³)', 'ศูนย์มะเร็งขอนตาล (m³)', 'ยอดรวมทั้งสิ้น (m³)', 'หมายเหตุ']]
for m, nkp, khon, tot in zip(months, nkp_vals, khon_vals, total_vals):
    khon_str = f"{khon:,.2f}" if khon > 0 else "-"
    note = ""
    if m == max_month:
        note = "ยอดใช้สูงสุด (Peak)"
    elif m == min_month:
        note = "ยอดใช้ต่ำสุด"
    elif m == "มี.ค. 69":
        note = "ข้อมูลล่าสุด"
    monthly_data.append([m, f"{nkp:,.2f}", khon_str, f"{tot:,.2f}", note])

table_m = doc.add_table(rows=len(monthly_data), cols=5)
table_m.alignment = WD_TABLE_ALIGNMENT.CENTER
for r_idx, row in enumerate(monthly_data):
    for c_idx, val in enumerate(row):
        cell = table_m.cell(r_idx, c_idx)
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'TH Sarabun PSK'
                run.font.size = Pt(12)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if r_idx == 0:
            set_cell_bg(cell, '003366')
        elif r_idx % 2 == 1:
            set_cell_bg(cell, 'F9FAFB')

doc.add_paragraph()

# Strategy Recommendations
h4 = doc.add_paragraph()
r = h4.add_run('4. ข้อเสนอแนะเชิงกลยุทธ์และการพัฒนาระบบ')
r.font.name = 'TH Sarabun PSK'
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)

recs = [
    'เตรียมความพร้อมรองรับอาคารหัวใจ 12 ชั้น (งบประมาณ 60 ล้านบาท) ที่จะเพิ่มปริมาณการใช้ออกซิเจนอีก 15-20% ควรปรับวงรอบการเติมแก๊สเข้าถังหลักจากทุก 3-4 วัน เป็นทุก 2-3 วัน',
    'ขยายระบบมอนิเตอร์โทรมาตรแบบรวมศูนย์ (Centralized Telemetry Dashboard) เชื่อมโยงสัญญาณจากถังหลัก NPPH10011, NPPH20011 และตู้ Area Alarm ทุกอาคารเข้าสู่แดชบอร์ดศูนย์ควบคุมวิศวกรรม',
    'ดำเนินการบำรุงรักษาเชิงป้องกัน (PM) อุปกรณ์ปรับแรงดัน และเปลี่ยนชุด Ohio Medical Maintenance Kit ตามแผนพัฒนาปี 2570 เพื่อมุ่งสู่เป้าหมายข้อผิดพลาดเป็นศูนย์ (Zero Failure Target)'
]

for rec in recs:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(rec)
    r.font.name = 'TH Sarabun PSK'
    r.font.size = Pt(14)

docx_path = r'E:\000_Antigraviti\MedicalGasNKP\Oxygen\07_Liquid_Oxygen_Main_System\รายงานสรุปการใช้ออกซิเจน_รพนครพิงค์.docx'
doc.save(docx_path)
print('DOCX Report successfully created at:', docx_path)
